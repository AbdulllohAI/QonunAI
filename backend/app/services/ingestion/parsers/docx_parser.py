"""DOCX parser — uses Word's own outline levels when the document has them."""
from __future__ import annotations

import io
import re

from app.services.ingestion.parsers.base import BaseParser, ParsedBlock, ParsedDocument
from app.services.ingestion.parsers.html_parser import _HEADING_RE
from app.services.lang.translit import normalize

_HEADING_STYLE_RE = re.compile(r"heading\s*(\d+)|заголовок\s*(\d+)|sarlavha\s*(\d+)", re.I)


class DocxParser(BaseParser):
    def parse(self, data: bytes | str, **kwargs) -> ParsedDocument:
        raw = data.encode() if isinstance(data, str) else data
        import docx

        document = docx.Document(io.BytesIO(raw))
        blocks: list[ParsedBlock] = []

        for paragraph in document.paragraphs:
            text = normalize(paragraph.text)
            if not text:
                continue

            level = _style_level(paragraph.style.name if paragraph.style else "")
            if level is None and _looks_bold_heading(paragraph) and len(text) < 250:
                level = 4
            if level is None and _HEADING_RE.match(text) and len(text) < 400:
                level = 4

            blocks.append(
                ParsedBlock(
                    text=text,
                    role="heading" if level is not None else "body",
                    level=level,
                )
            )

        # Tables carry schedules, tariffs and penalty scales — worth keeping.
        for table in document.tables:
            for row in table.rows:
                cells = [normalize(c.text) for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    blocks.append(ParsedBlock(text=line, role="table"))

        core = document.core_properties
        title = normalize(core.title) if core and core.title else None
        if not title and blocks:
            title = blocks[0].text[:300]

        return ParsedDocument(blocks=blocks, title=title, meta={"format": "docx"})


def _style_level(style_name: str) -> int | None:
    match = _HEADING_STYLE_RE.search(style_name or "")
    if not match:
        return None
    for group in match.groups():
        if group:
            return int(group)
    return None


def _looks_bold_heading(paragraph) -> bool:
    runs = [r for r in paragraph.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)
